from docx import Document

document = Document()

document.add_heading('Pizzas', 0)

pizza_flavours = ["Green", "Yellow", "Triple Cheese & Garlic with Cranberry Swirl"]

for pizza in pizza_flavours:
	document.add_paragraph(
	    pizza
	)


document.save('dem2o.docx')
